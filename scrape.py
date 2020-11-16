from chords import chords  # Importing list of chords from UltimateGuitar

import argparse
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor
import os
import re
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Preformatted, Spacer, Paragraph
from selenium import webdriver


def get_site_content(url):
    """Scraping the given URL for chords and lyrics and then making PDF and/or word document from it."""

    # Using a driver to be able to scrape from dynamic <pre>
    driver = webdriver.Chrome(executable_path="chromedriver.exe")
    driver.get(url)
    html = driver.page_source
    driver.quit()

    # Start making soup
    soup = BeautifulSoup(html, features="lxml")
    song = soup.find('pre')  # Find the chords and lyrics from dynamic <pre> tag

    return song.get_text().splitlines()


def make_pdf(artist, title, site_content):
    """Taking the scraped site content from the 'get_site_content'-method and create PDF using Reportlab."""

    pdf_directory = './pdfs/'

    if not os.path.exists(pdf_directory):
        os.makedirs(pdf_directory)

    styles = getSampleStyleSheet()

    # Style for lyrics
    styles.add(ParagraphStyle(name='Lyrics',
                              parent=styles['Code'],
                              fontName='Courier',
                              fontSize=10,
                              leading=12))
    style_lyrics = styles['Lyrics']

    # Style for Chords
    styles.add(ParagraphStyle(name='Chords',
                              parent=styles['Code'],
                              fontName='Courier-Bold',
                              fontSize=10,
                              leading=12,
                              textColor=colors.HexColor("#C00000")))
    style_chords = styles['Chords']

    # Style for SongTitle
    styles.add(ParagraphStyle(name='SongTitle',
                              parent=styles['Heading1'],
                              fontName='Courier-Bold',
                              leading=10,
                              leftIndent=35,
                              textColor=colors.HexColor("#002060")))
    style_title = styles['SongTitle']

    # Style for Artist
    styles.add(ParagraphStyle(name='Artist',
                              parent=styles['Heading4'],
                              fontName='Courier',
                              leftIndent=35))
    style_artist = styles['Artist']

    pdf = SimpleDocTemplate(pdf_directory + '{} - {}.pdf'.format(artist.lower(), title.lower()),
                            rightMargin=35,
                            leftMargin=0,
                            topMargin=35,
                            bottomMargin=35)

    # List of Flowables to be rendered to the PDF document
    text = [Paragraph(title.upper(), style_title), Paragraph(artist.title(), style_artist)]
    text.append(Spacer(1, 10))

    chords_first = False  # Make sure cords is the first to be added to the PDF after title and artist, not any nonsence.

    # If chord apply chord style, if lyrics apply lyric style if blank line add blank line
    # THIS LOOKS AWFUL
    for line in site_content:
        if not line and chords_first:
            text.append(Spacer(1, 10))
        if is_chord(line) and line != '':
            chords_first = True
            text.append(Preformatted(line, style_chords))
        else:
            if chords_first:
                text.append(Preformatted(line, style_lyrics))

    pdf.build(text)


def make_doc(artist, title, site_content):
    """Taking the scraped site content from the 'get_site_content'-method and create Word doc using Docx module."""

    doc_directory = './docs/'

    if not os.path.exists(doc_directory):
        os.makedirs(doc_directory)

    document = Document()

    margins = document.sections
    for side in margins:
        side.top_margin = Cm(0.1)
        side.bottom_margin = Cm(0.1)
        side.left_margin = Cm(1.5)
        side.right_margin = Cm(1.5)

    styles = document.styles

    # Adding Title and styling
    doc_title = document.add_paragraph()
    doc_title.add_run(title.upper()).bold = True
    doc_title_style = document.styles['No Spacing']
    doc_title_font = doc_title_style.font
    doc_title_font.color.rgb = RGBColor(0x00, 0x20, 0x60)
    doc_title_font.name = 'Courier'
    doc_title_font.size = Pt(20)
    doc_title.style = document.styles['No Spacing']

    # Adding Artist and style
    doc_artist = document.add_paragraph(artist.title())
    doc_artist.add_run('\n')
    doc_artist_style = document.styles['Normal']
    doc_artist_font = doc_artist_style.font
    doc_artist_font.name = 'Courier'
    doc_artist_font.size = Pt(10)
    doc_artist.style = document.styles['Normal']

    # Make sure cords is the first to be added to the PDF after title and artist.
    chords_first = False

    # Style chordline
    chord_line_style = styles.add_style('Chordline', WD_STYLE_TYPE.PARAGRAPH)
    chord_line_style.base_style = styles['No Spacing']
    chord_line_font = chord_line_style.font
    chord_line_font.name = 'Courier'
    chord_line_font.size = Pt(11)
    chord_line_font.color.rgb = RGBColor(0xc0, 0x00, 0x00)

    # Style lyrics
    lyric_line_style = styles.add_style('Lyrics', WD_STYLE_TYPE.PARAGRAPH)
    lyric_line_style.base_style = styles['No Spacing']
    lyric_line_font = lyric_line_style.font
    lyric_line_font.name = 'Courier'
    lyric_line_font.size = Pt(11)
    lyric_line_font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Adding lyrics and chords + styling
    for line in site_content:
        if is_chord(line) and line != '':
            chords_first = True
            chord_line = document.add_paragraph()
            chord_line.add_run(line).bold = True
            chord_line.style = document.styles['Chordline']

        else:
            if chords_first:
                lyric_line = document.add_paragraph(line)
                lyric_line.style = document.styles['Lyrics']

    document.save(doc_directory + '{} - {}.docx'.format(artist.lower(), title.lower()))  # Save document


def is_chord(string):
    """ Check if the supplied string (line from Ultimate Guitar) is chords or lyrics."""
    # Compare each line with chords to see if lyrics or chords
    # VERY in efficient! Hacked together to move on with project... Will need tuning.
    line_as_list = string.split(' ')

    for chord in chords:
        if chord in line_as_list and not any('-' in x for x in line_as_list):
            return True

# Parser to handle URL in argument.
parser = argparse.ArgumentParser()
parser.add_argument('Artist', type=str, help='Artist name. Ex. "Bruce Springsteen"')
parser.add_argument('Title', type=str, help='Title name. Ex. "Kingdom of days"')
parser.add_argument('URL', type=str, help='URL for song. Ex. "FULL URL TO CHORDS"')
parser.add_argument('-p', action='store_true', help='PDF document to be made')
parser.add_argument('-w', action='store_true', help='Word document to be made')
parser.add_argument('-b', action='store_true', help='Both PDF and .docx to be made')

if __name__ == "__main__":
    args = parser.parse_args()
    # Argument checking for correct flags
    if not (args.p or args.w or args.b):
        parser.error('No flag to determine document type. Use -p for PDF, -w for Word document or -b for both.')
    if args.p and args.w:
        parser.error('No need to use both -p and -w, use -b instead to create both document types.')
    if args.b and (args.p or args.w):
        parser.error('Do not use -b together with -p or -w. Use only -w, -p or -b.')

    # Getting the content from ultimateguitar.
    site_content = get_site_content(args.URL)

    # Creating the documents according to flags in arguments.
    if args.b:
        make_pdf(args.Artist, args.Title, site_content)
        make_doc(args.Artist, args.Title, site_content)
    if args.w:
        make_doc(args.Artist, args.Title, site_content)
    if args.p:
        make_pdf(args.Artist, args.Title, site_content)
